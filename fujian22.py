from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.enum.table import WD_TABLE_ALIGNMENT,WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor
from docx.oxml.ns import qn
from docx.shared import Cm,Inches
import pymysql

unit="市水利局"

def fujian2(unit):
    db = pymysql.connect(host = '192.168.75.128', port=3306, user = 'zhangbo', password = '22', database = 'testdb')
    cur = db.cursor()
    #sql = """select * from 3k where unit='潍坊市交通运输局' and host='10.80.13.62' order by system,host"""
    sql = """select * from 3k where unit='{}' order by system,host""".format(unit)
    cur.execute(sql)
    data = cur.fetchall()
    cur.close()
    db.close()

    rows=len(data)+4
    # 创建一个新的Word文档
    doc = Document()

    head = doc.add_heading(level=3)
    run = head.add_run('附件2')
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run.font.size = Pt(16)
    run.font.bold = False
    p = doc.add_paragraph()

    head = doc.add_heading(level=3)
    run1 = head.add_run('漏洞处置情况反馈表')
    run1.font.color.rgb = RGBColor(0, 0, 0)
    run1.font.name = '仿宋'
    run1._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
    run1.font.size = Pt(20)
    run1.font.bold = False
    head.alignment=WD_TABLE_ALIGNMENT.CENTER

    # col1 = table.columns[0]
    # col1.width = Inches(1.3)
    # 添加一个4列10行的表格
    table = doc.add_table(rows=rows, cols=4,style="Table Grid")
    table.autofit = False
    table.allow_autofit = False
    table.columns[0].width = Inches(1.3)
    table.columns[1].width = Inches(2.6)
    table.columns[2].width = Inches(1.3)
    table.columns[3].width = Inches(1.3)

    # 设置第一行的字体为仿宋14加粗
    for cell in table.rows[0].cells:
        run = cell.paragraphs[0].add_run()
        run.font.name = u'仿宋'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋')
        run.font.size = Pt(14)
        run.font.bold = True

    # 设置第二行到第九行的字体为仿宋11
    for row in table.rows[1:rows-2]:
        for cell in row.cells:
            run = cell.paragraphs[0].add_run()
            run.font.name = 'FangSong'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
            run.font.size = Pt(11)
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 合并最后一行的4个单元格，并设置字体为仿宋14
    merged_cell = table.cell(rows-1, 0).merge(table.cell(rows-1, 3))
    run = merged_cell.paragraphs[0].add_run()
    run.font.name = 'FangSong'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
    run.font.size = Pt(14)

    border_xml = """
    <w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
      <w:top w:val="single"/>
      <w:left w:val="single"/>
      <w:bottom w:val="single"/>
      <w:right w:val="single"/>
    </w:tcBorders>
    """
    for row in table.rows[0:]:
        for cell in row.cells:
            cell._tc.get_or_add_tcPr().append(parse_xml(border_xml))

    table.cell(0, 0).text = "IP地址"
    table.cell(0, 1).text = "漏洞名称"
    table.cell(0, 2).text = "是否整改"
    table.cell(0, 3).text = "未整改原因"
    table.cell(rows-3, 0).text = "*负责人"
    table.cell(rows-3, 1).text = "\n\n"
    table.cell(rows-2, 1).text = "\n\n"
    table.cell(rows-3, 2).text = "*联系电话"
    table.cell(rows-2, 0).text = "*处置人"
    table.cell(rows-2, 2).text = "*联系电话"
    table.cell(rows-1, 0).text = "\n\n*网络安全负责人：（签字）\n\n"
    table.cell(rows-1, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    table.cell(rows-1, 0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table.cell(rows-1, 0).paragraphs[0].runs[0].font.size = Pt(14)
    table.cell(rows-1, 0).paragraphs[0].runs[0].font.name = u'仿宋'
    table.cell(rows-1, 0).paragraphs[0].runs[0].font.bold = False
    table.cell(rows-1, 0).paragraphs[0].runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋')


    for i in range(0,4):
        table.cell(0, i).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(0, i).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        table.cell(0, i).paragraphs[0].runs[0].font.size = Pt(14)
        table.cell(0, i).paragraphs[0].runs[0].font.name = u'仿宋'
        table.cell(0, i).paragraphs[0].runs[0].font.bold = True
        table.cell(0, i).paragraphs[0].runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋')

    for i in (0,2):
        for j in (1,2):
            table.cell(rows-1-j, i).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            table.cell(rows-1-j, i).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            table.cell(rows-1-j, i).paragraphs[0].runs[0].font.size = Pt(14)
            table.cell(rows-1-j, i).paragraphs[0].runs[0].font.name = u'仿宋'
            table.cell(rows-1-j, i).paragraphs[0].runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋')
            if i==0:
                table.cell(rows-1-j, i).paragraphs[0].runs[0].font.bold = True
            else:
                pass

    head = doc.add_heading(level=3)
    run1 = head.add_run('注：*为必填项')
    run1.font.color.rgb = RGBColor(0, 0, 0)
    run1.font.name = u'仿宋'
    run1._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋')
    run1.font.size = Pt(14)
    run1.font.bold = False
    # 保存文档

    for i in range(len(data)):
        print(data[i])

    for row in range(1,rows-3):
        table.cell(row, 0).text = data[row-1][3]
        table.cell(row, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(row, 0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        table.cell(row, 0).paragraphs[0].runs[0].font.size = Pt(11)
        table.cell(row, 0).paragraphs[0].runs[0].font.name = u'仿宋'
        table.cell(row, 0).paragraphs[0].runs[0].font.bold = False
        table.cell(row, 0).paragraphs[0].runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋')
        table.cell(row, 1).text = data[row-1][5]
        table.cell(row, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(row, 1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        table.cell(row, 1).paragraphs[0].runs[0].font.size = Pt(11)
        table.cell(row, 1).paragraphs[0].runs[0].font.name = u'仿宋'
        table.cell(row, 1).paragraphs[0].runs[0].font.bold = False
        table.cell(row, 1).paragraphs[0].runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋')
        table.cell(row, 3).text = "\n\n"

    doc.save('table.docx')

fujian2(unit)