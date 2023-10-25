from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import RGBColor
from docx.oxml.ns import qn
import pymysql



rows=6
# 创建一个新的Word文档
doc = Document()

head = doc.add_heading(level=3)
run = head.add_run('附件2')
run.font.color.rgb = RGBColor(0, 0, 0)
run.font.name = '黑体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
run.font.size = Pt(16)



p = doc.add_paragraph()
# 添加一个4列17行的表格
table = doc.add_table(rows=rows, cols=4)

# 合并第一行的所有列
cell = table.cell(0, 0)
cell.merge(table.cell(0, 1))
cell.merge(table.cell(0, 2))
cell.merge(table.cell(0, 3))

# 设置第一行的内容
cell.text = "漏洞处置情况反馈表"
run = cell.paragraphs[0].runs[0]
run.font.size = Pt(12)
run.font.name = '仿宋'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
run.font.size = Pt(20)
cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# 设置第二行的内容
table.cell(1, 0).text = "IP地址"
table.cell(1, 1).text = "漏洞名称"
table.cell(1, 2).text = "是否整改"
table.cell(1, 3).text = "未整改原因"

# 设置第三行到第十行的第一列的内容



for i in range(2, rows):
    table.cell(i, 0).text = str(i+1)

# 设置所有单元格的文字居中
for row in table.rows:
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_TABLE_ALIGNMENT.CENTER

# 设置第二行开始的单元格为实线
border_xml = """
<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:top w:val="single"/>
  <w:left w:val="single"/>
  <w:bottom w:val="single"/>
  <w:right w:val="single"/>
</w:tcBorders>
"""
for row in table.rows[1:]:
    for cell in row.cells:
        cell._tc.get_or_add_tcPr().append(parse_xml(border_xml))



# 保存文档
# doc.save("table.docx")