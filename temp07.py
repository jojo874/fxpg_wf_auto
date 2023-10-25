import os
import sys
import time
from docx import Document
import docx
from docx.shared import RGBColor
import pymysql
from docx.shared import Pt,Cm,Inches
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT,WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH


def create_doc_table(data):
    doc_name = "2023_2.docx"
    doc = docx.Document(doc_name)

    num_china = ['一、 ' , '二、 ' , '三、 ' , '四、 ' , '五、 ' , '六、 ' , '七、 ' , '八、 ' , '九、 ' , '十、 ' , '十一、 ' , '十二、 ' , '十三、 ' , '十四、 ' , '十五、 ' , '十六、 ' , '十七、 ' , '十八、 ' , '十九、 ' , '二十、 ' , '二十一、 ']
    for i in data:

        if i[4] != '':
            text_str = num_china.pop(0) + i[4] + "系统\n\n"
            text_str = text_str + "资产 "+i[3]
            paragraph_content = doc.add_paragraph()
            run = paragraph_content.add_run(text_str)
            run.font.name = '仿宋'
            run.font.bold = False
            run.font.size = Pt(16)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')

        else:
            if i[3] != '':
                text_str = ''
                text_str = text_str + "资产 "+i[3]
                paragraph_content = doc.add_paragraph()
                run = paragraph_content.add_run(text_str)
                run.font.name = '仿宋'
                run.font.bold = False
                run.font.size = Pt(16)
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            else:
                # paragraph_content = doc.add_paragraph()
                # run = paragraph_content.add_run(text_str)
                # run.font.name = u'宋体'
                # run.font.bold = True
                # run.font.size = Pt(16)
                pass



        field_list = ["漏洞名称", i[5]]
        if "CVE" in i[5]:
            content_list = [["风险等级：", "高危"], ["CVE编号：", i[5].split('(')[-1][:-1]], ["端口(服务)：", "-"],["风险描述：", i[6]],["危害影响：", i[6]],["解决方案：", i[7]],["参考资料：", "https://www.cnnvd.org.cn/"]]
        else:
            content_list = [["风险等级：", "高危"], ["CVE编号：", "-"], ["端口(服务)：", "-"],["风险描述：", i[6]],["危害影响：", i[6]],["解决方案：", i[7]],["参考资料：", "https://www.cnnvd.org.cn/"]]
        table = doc.add_table(rows=8, cols=len(field_list),style="Table Grid")
        table.autofit = False
        table.allow_autofit = False

        col1=table.columns[0]
        col1.width = Inches(1.3)
        #col1.width = Pt(20)
        col2=table.columns[1]
        col2.width = Inches(4.6)
        #col2.width = Pt(80)



        #table.style = "Light List Accent 1"
        cells = table.rows[0].cells
        for i in range(len(field_list)):
            new = field_list[i]
            cells[i].text = ""
            paragraph = cells[i].paragraphs[0]
            #paragraph.add_run(new).bold = True
            a=paragraph.add_run(new)
            a.font.name = '仿宋'
            a.font.bold = True
            a.font.size = Pt(11)
            a._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋')
        shading_elm_1 = parse_xml(r'<w:shd {} w:fill="BEBEBE"/>'.format(nsdecls('w')))
        table.rows[0].cells[0]._tc.get_or_add_tcPr().append(shading_elm_1)
        shading_elm_2 = parse_xml(r'<w:shd {} w:fill="BEBEBE"/>'.format(nsdecls('w')))
        table.rows[0].cells[1]._tc.get_or_add_tcPr().append(shading_elm_2)


        for rows_id in range(7):
            if rows_id == 1:
                item = content_list[rows_id]
                print(item)
                row_cells_data = table.rows[rows_id + 1].cells
                for i in range(len(field_list)):
                    #row_cells_data[i].text = item[i]
                    new1 = item[i]
                    row_cells_data[i].text = ""
                    paragraph = row_cells_data[i].paragraphs[0]
                    a=paragraph.add_run(new1)
                    a.font.name = u'仿宋'
                    a.font.bold = True
                    a.font.size = Pt(11)
                    a._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋')
            else:
                item = content_list[rows_id]
                print(item)
                row_cells_data = table.rows[rows_id + 1].cells
                for i in range(len(field_list)):
                    # row_cells_data[i].text = item[i]
                    new1 = item[i]
                    row_cells_data[i].text = ""
                    paragraph = row_cells_data[i].paragraphs[0]
                    a = paragraph.add_run(new1)
                    a.font.name = u'仿宋'
                    a.font.bold = False
                    a.font.size = Pt(11)
                    a._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋')

        text_str2 = "\n"
        paragraph_content2 = doc.add_paragraph()
        run2 = paragraph_content2.add_run(text_str2)
        run2.font.name = u'宋体'
        run2.font.bold = True
        run2.font.size = Pt(11)


    time_str = str(time.time())
    print("time_str====", time_str)
    db = pymysql.connect(host = '192.168.75.128', port=3306, user = 'zhangbo', password = '22', database = 'testdb')
    cur = db.cursor()
    #sql = """select * from 3k where unit='潍坊市交通运输局' and host='10.80.13.62' order by system,host"""
    sql = """select count(unit) from 3k where unit='{}'""".format(data[0][1])
    cur.execute(sql)
    count = cur.fetchone()
    cur.close()
    db.close()

#确定市,区,县,单位
    attrList_shi = ["高密市", "青州市", "诸城市", "寿光市", "安丘市", "昌邑市"]
    attrList_qu = ["奎文区", "潍城区", "房子区", "寒亭区", "高新区", "滨海区", "保税区", "峡山区", "经济开发区"]
    attrList_xian = ["临朐县", "昌乐县"]
    attr = "单位"
    for i in attrList_shi:
        if i in data[0][0]:
            attr = "市"
    for i in attrList_qu:
        if i in data[0][0]:
            attr = "区"
    for i in attrList_xian:
        if i in data[0][0]:
            attr = "县"

    params = {
        "x": count[0],
        "unit": data[0][1],
        "att": attr,
    }
    replace_placeholder(doc, params)

#添加分页符
    doc.add_page_break()
##************************************************************************************************
##附件2
    db = pymysql.connect(host = '192.168.75.128', port=3306, user = 'zhangbo', password = '22', database = 'testdb')
    cur = db.cursor()
    #sql = """select * from 3k where unit='潍坊市交通运输局' and host='10.80.13.62' order by system,host"""
    sql = """select * from 3k where unit='{}' order by system,host""".format(data[0][0])
    cur.execute(sql)
    data = cur.fetchall()
    cur.close()
    db.close()

    rows=len(data)+4
    # 创建一个新的Word文档

    head = doc.add_heading(level=3)
    run = head.add_run('附件2')
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run.font.size = Pt(16)
    run.font.bold = False
    p = doc.add_paragraph()

    head = doc.add_heading(level=3)
    run1 = head.add_run('漏洞处置情况反馈表')
    run1.font.color.rgb = RGBColor(0, 0, 0)
    run1.font.name = '仿宋'
    run1._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
    run1.font.size = Pt(20)
    run1.font.bold = False
    head.alignment=WD_TABLE_ALIGNMENT.CENTER

    # col1 = table.columns[0]
    # col1.width = Inches(1.3)
    # 添加一个4列10行的表格
    table = doc.add_table(rows=rows, cols=4,style="Table Grid")
    table.autofit = False
    table.allow_autofit = False
    table.columns[0].width = Inches(1.3)
    table.columns[1].width = Inches(2.6)
    table.columns[2].width = Inches(1.3)
    table.columns[3].width = Inches(1.3)

    # 设置第一行的字体为仿宋14加粗
    for cell in table.rows[0].cells:
        run = cell.paragraphs[0].add_run()
        run.font.name = u'仿宋'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋')
        run.font.size = Pt(14)
        run.font.bold = True

    # 设置第二行到第九行的字体为仿宋11
    for row in table.rows[1:rows-2]:
        for cell in row.cells:
            run = cell.paragraphs[0].add_run()
            run.font.name = 'FangSong'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
            run.font.size = Pt(11)
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 合并最后一行的4个单元格，并设置字体为仿宋14
    merged_cell = table.cell(rows-1, 0).merge(table.cell(rows-1, 3))
    run = merged_cell.paragraphs[0].add_run()
    run.font.name = 'FangSong'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
    run.font.size = Pt(14)

    border_xml = """
    <w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
      <w:top w:val="single"/>
      <w:left w:val="single"/>
      <w:bottom w:val="single"/>
      <w:right w:val="single"/>
    </w:tcBorders>
    """
    for row in table.rows[0:]:
        for cell in row.cells:
            cell._tc.get_or_add_tcPr().append(parse_xml(border_xml))

    table.cell(0, 0).text = "IP地址"
    table.cell(0, 1).text = "漏洞名称"
    table.cell(0, 2).text = "是否整改"
    table.cell(0, 3).text = "未整改原因"
    table.cell(rows-3, 0).text = "*负责人"
    table.cell(rows-3, 1).text = "\n\n"
    table.cell(rows-2, 1).text = "\n\n"
    table.cell(rows-3, 2).text = "*联系电话"
    table.cell(rows-2, 0).text = "*处置人"
    table.cell(rows-2, 2).text = "*联系电话"
    table.cell(rows-1, 0).text = "\n*网络安全负责人：（签字）\n"
    table.cell(rows-1, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    table.cell(rows-1, 0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table.cell(rows-1, 0).paragraphs[0].runs[0].font.size = Pt(14)
    table.cell(rows-1, 0).paragraphs[0].runs[0].font.name = u'仿宋'
    table.cell(rows-1, 0).paragraphs[0].runs[0].font.bold = False
    table.cell(rows-1, 0).paragraphs[0].runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋')


    for i in range(0,4):
        table.cell(0, i).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(0, i).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        table.cell(0, i).paragraphs[0].runs[0].font.size = Pt(14)
        table.cell(0, i).paragraphs[0].runs[0].font.name = u'仿宋'
        table.cell(0, i).paragraphs[0].runs[0].font.bold = True
        table.cell(0, i).paragraphs[0].runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋')

    for i in (0,2):
        for j in (1,2):
            table.cell(rows-1-j, i).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            table.cell(rows-1-j, i).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            table.cell(rows-1-j, i).paragraphs[0].runs[0].font.size = Pt(14)
            table.cell(rows-1-j, i).paragraphs[0].runs[0].font.name = u'仿宋'
            table.cell(rows-1-j, i).paragraphs[0].runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋')
            if i==0:
                table.cell(rows-1-j, i).paragraphs[0].runs[0].font.bold = True
            else:
                pass

    head = doc.add_heading(level=3)
    run1 = head.add_run('注：*为必填项')
    run1.font.color.rgb = RGBColor(0, 0, 0)
    run1.font.name = u'仿宋'
    run1._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋')
    run1.font.size = Pt(14)
    run1.font.bold = False
    # 保存文档

    # for i in range(len(data)):
    #     print(data[i])

    for row in range(1,rows-3):
        table.cell(row, 0).text = data[row-1][3]
        table.cell(row, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(row, 0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        table.cell(row, 0).paragraphs[0].runs[0].font.size = Pt(11)
        table.cell(row, 0).paragraphs[0].runs[0].font.name = u'仿宋'
        table.cell(row, 0).paragraphs[0].runs[0].font.bold = False
        table.cell(row, 0).paragraphs[0].runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋')
        table.cell(row, 1).text = data[row-1][5]
        table.cell(row, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(row, 1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        table.cell(row, 1).paragraphs[0].runs[0].font.size = Pt(11)
        table.cell(row, 1).paragraphs[0].runs[0].font.name = u'仿宋'
        table.cell(row, 1).paragraphs[0].runs[0].font.bold = False
        table.cell(row, 1).paragraphs[0].runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋')
        table.cell(row, 3).text = "\n"

    new_doc_name = str(count[0]) + data[0][1] + "安全检查结果通知.docx"
    doc.save(new_doc_name)


#create_doc_table()

def getFromMysql(unit):
    db = pymysql.connect(host = '192.168.75.128', port=3306, user = 'zhangbo', password = '22', database = 'testdb')
    cur = db.cursor()
    #sql = """select * from 3k where unit='潍坊市交通运输局' and host='10.80.13.62' order by system,host"""
    sql = """select * from 3k where unit='{}' order by system,host""".format(unit)
    cur.execute(sql)
    data = cur.fetchall()
    cur.close()
    db.close()

    system = data[0][4]
    host = data[0][3]
    system = ''
    host = ''

    list_all = []
    for i in data:
        lst = list(i)
        if lst[4] != system:
            if lst[3] != host:
                host = lst[3]
                #print(lst)
                system = lst[4]
                list_all.append(lst)
                #print(type(lst))
            else:
                lst[3] = ''
                system = lst[4]
                list_all.append(lst)
        else:
            if lst[3] != host:
                host = lst[3]
                lst[4] = ''
                #print(lst[4])
                list_all.append(lst)
            else:
                lst[3] = ''
                lst[4] = ''
                # print(lst[4])
                list_all.append(lst)
    return list_all

#替换关键字
def replace_placeholder(doc, params):
    """替换占位符"""
    for paragraph in doc.paragraphs:
        for param in params:
            pv = str(params[param])
            ph = f'ph_{param}'
            if ph in paragraph.text:
                for run in paragraph.runs:
                    if ph in run.text:
                        run.text = run.text.replace(ph, pv)
                        run.italic = False
# unit='潍坊市交通运输局'
#
# create_doc_table(getFromMysql(unit))

def getUnit():
    db = pymysql.connect(host = '192.168.75.128', port=3306, user = 'zhangbo', password = '22', database = 'testdb')
    cur = db.cursor()
    sql = """select unit from testdb.3k group by unit"""
    cur.execute(sql)
    data = cur.fetchall()
    cur.close()
    db.close()
    return data

units = getUnit()
for u in units:
    print(u[0])
    create_doc_table(getFromMysql(u[0]))
